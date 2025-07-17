import streamlit as st
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from fpdf import FPDF
import io
from PIL import Image
import img2pdf
from docx2pdf import convert
import os
import tempfile
import pythoncom

st.set_page_config(page_title="File Converter",layout="wide", page_icon="🔄")

def pdf_to_word(uploaded_file):
    pdf = PdfReader(uploaded_file)
    doc = Document()
    
    for page in pdf.pages:
        text = page.extract_text()
        if text:
            doc.add_paragraph(text)
    
    word_io = io.BytesIO()
    doc.save(word_io)
    word_io.seek(0)
    return word_io

def word_to_pdf(uploaded_file):
    # Create a temporary file to save the uploaded .docx
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_word:
        temp_word.write(uploaded_file.read())
        temp_word_path = temp_word.name  # Get temp file path

    # Create a temporary PDF file
    temp_pdf_path = temp_word_path.replace(".docx", ".pdf")

    # Initialize COM for Windows before conversion
    pythoncom.CoInitialize()
    convert(temp_word_path, temp_pdf_path)
    pythoncom.CoUninitialize()

    # Read the generated PDF into memory
    with open(temp_pdf_path, "rb") as f:
        pdf_io = io.BytesIO(f.read())

    # Clean up temporary files
    os.remove(temp_word_path)
    os.remove(temp_pdf_path)

    return pdf_io

def ppt_to_pdf(uploaded_file):
    import comtypes.client

    # Save uploaded ppt/pptx to a temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pptx") as temp_ppt:
        temp_ppt.write(uploaded_file.read())
        temp_ppt_path = temp_ppt.name

    # Output PDF path
    temp_pdf_path = temp_ppt_path.replace(".pptx", ".pdf")

    # COM init for Windows
    pythoncom.CoInitialize()
    powerpoint = comtypes.client.CreateObject("PowerPoint.Application")
    powerpoint.Visible = 1

    try:
        presentation = powerpoint.Presentations.Open(temp_ppt_path, WithWindow=False)
        presentation.SaveAs(temp_pdf_path, 32)  # 32 = PDF
        presentation.Close()
    finally:
        powerpoint.Quit()
        pythoncom.CoUninitialize()

    # Read PDF as BytesIO
    with open(temp_pdf_path, "rb") as f:
        pdf_io = io.BytesIO(f.read())

    # Cleanup
    os.remove(temp_ppt_path)
    os.remove(temp_pdf_path)

    return pdf_io


def merge_pdfs(uploaded_files):
    merger = PdfWriter()
    for pdf in uploaded_files:
        merger.append(pdf)
    merged_io = io.BytesIO()
    merger.write(merged_io)
    merger.close()
    merged_io.seek(0)
    return merged_io

def merge_word_files(uploaded_files):
    merged_doc = Document(uploaded_files[0])  # Start with the first document
    
    for file in uploaded_files[1:]:
        doc = Document(file)
        
        for element in doc.element.body:
            merged_doc.element.body.append(element)  # Append elements directly to preserve styling
    
    merged_io = io.BytesIO()
    merged_doc.save(merged_io)
    merged_io.seek(0)
    
    return merged_io

def images_to_pdf(uploaded_images):
    img_list = []
    
    for image_file in uploaded_images:
        # Open image
        image = Image.open(image_file)
        
        # Convert to RGB (img2pdf does not support RGBA or P modes)
        if image.mode in ("RGBA", "P"):
            image = image.convert("RGB")
        
        # Save image to BytesIO
        img_io = io.BytesIO()
        image.save(img_io, format="JPEG")
        img_io.seek(0)
        
        img_list.append(img_io.getvalue())  # Append image as bytes
    
    # Convert images to a single PDF
    pdf_bytes = img2pdf.convert(img_list)
    
    # Store PDF in BytesIO
    pdf_io = io.BytesIO()
    pdf_io.write(pdf_bytes)
    pdf_io.seek(0)

    return pdf_io




st.title("File Converter")
tabs = st.tabs(["PDF to Word", "Word to PDF", "PPT to PDF", "Merge PDFs", "Merge Word Files","Images to PDF"])

with tabs[0]:
    st.header("Convert PDF to Word")
    pdf_file = st.file_uploader("Upload a PDF file", type=["pdf"], key="pdf_to_word")
    if pdf_file:
        word_io = pdf_to_word(pdf_file)
        st.download_button("Download Word File", word_io, "Converted.docx")

with tabs[1]:
    st.header("Convert Word to PDF")
    word_file = st.file_uploader("Upload a Word file", type=["docx"], key="word_to_pdf")
    if word_file:
        pdf_io = word_to_pdf(word_file)
        st.download_button("Download PDF File", pdf_io, "Converted.pdf")

with tabs[2]:
    st.header("Convert PPT to PDF")
    ppt_file = st.file_uploader("Upload a PowerPoint file", type=["ppt", "pptx"], key="ppt_to_pdf")
    if ppt_file:
        pdf_io = ppt_to_pdf(ppt_file)
        st.download_button("Download PDF File", pdf_io, "Presentation.pdf")

with tabs[3]:
    st.header("Merge PDF Files")
    pdf_files = st.file_uploader("Upload multiple PDFs", type=["pdf"], accept_multiple_files=True, key="merge_pdfs")
    if pdf_files:
        merged_pdf = merge_pdfs(pdf_files)
        st.download_button("Download Merged PDF", merged_pdf, "Merged.pdf")

with tabs[4]:
    st.header("Merge Word Files")
    word_files = st.file_uploader("Upload multiple Word files", type=["docx"], accept_multiple_files=True, key="merge_word")
    if word_files:
        merged_word = merge_word_files(word_files)
        st.download_button("Download Merged Word File", merged_word, "Merged.docx")

with tabs[5]:
    st.header("Convert Images to PDF")
    image_files = st.file_uploader("Upload multiple images", type=["jpg", "jpeg", "png"], accept_multiple_files=True, key="images_to_pdf")
    if image_files:
        pdf_io = images_to_pdf(image_files)
        st.download_button("Download PDF File", pdf_io, "Images.pdf")       
